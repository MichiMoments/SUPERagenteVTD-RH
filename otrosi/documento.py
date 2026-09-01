"""Generación determinística de un otrosí a partir de su tipo y del payload.

Sin dependencia de Streamlit ni de ninguna librería de plantillas: el cuerpo del
tipo se sustituye con una regex contra un diccionario de claves declaradas, así
que el texto que escriba una persona en la web no puede evaluar nada.
"""

import io
import re
import unicodedata
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Emu, Inches, Pt

from . import RAIZ

_PLANTILLAS = Path(__file__).resolve().parent / "plantillas"

LOGO = _PLANTILLAS / "LogoUninades.png"

TITULO = "OTROSÍ AL CONTRATO DE TRABAJO"

PIE = (
    "Dirección de Gestión Humana y Desarrollo Organizacional",
    "Carrera 1 No. 18A-81, Edificio Monjas Bogotá – Colombia. Tels.: [571] 3394949/99 "
    "Ext.: 3884 - 2243 Línea directa: [571] 3324374",
    "Universidad de los Andes | Vigilada Mineducación. Reconocimiento como Universidad: "
    "Decreto 1297 del 30 de mayo de 1964.",
    "Reconocimiento personería jurídica: Resolución 28 del 23 de febrero de 1949 Minjusticia",
)

MESES = (
    "enero",
    "febrero",
    "marzo",
    "abril",
    "mayo",
    "junio",
    "julio",
    "agosto",
    "septiembre",
    "octubre",
    "noviembre",
    "diciembre",
)

def fecha_larga(valor):
    """date/datetime -> '3 de agosto de 2026'."""
    if valor is None:
        return ""
    return f"{valor.day} de {MESES[valor.month - 1]} de {valor.year}"


def cedula(valor):
    """52832252 -> '52.832.252'. Tolera '52.832.252' de una carga masiva."""
    digitos = re.sub(r"\D", "", str(valor or ""))
    if not digitos:
        return ""
    return f"{int(digitos):,}".replace(",", ".")


def mayuscula_inicial(texto):
    """'la Teletrabajadora' -> 'La Teletrabajadora'."""
    # `|capitalize` de Jinja no sirve aquí: baja el resto de la cadena a minúsculas.
    return texto[:1].upper() + texto[1:]


# Lista fija y cerrada. No hay filtros de formato porque el formato lo decide el tipo
# del campo: {{fecha_ingreso}} imprime «3 de agosto de 2026» por ser de tipo fecha.
FILTROS = {
    "mayuscula": mayuscula_inicial,
    "mayusculas": str.upper,
    "minusculas": str.lower,
}

# Cómo se imprime cada tipo de campo; texto y lista salen tal cual.
FORMATOS = {"cedula": cedula, "fecha": fecha_larga}

# El separador de filtro es «:» y no «|»: el «|» delimita las celdas de tabla del
# conversor y el bloque de firmas lleva un marcador dentro de una celda.
MARCADOR = re.compile(r"\{\{\s*([^{}]*?)\s*\}\}")

_SEPARADORES = (":", "|")  # el «|» se tolera por si alguien viene de Jinja


def separar_marcador(interior):
    """Interior de un {{…}} -> (clave, filtro): 'x : mayuscula' -> ('x', 'mayuscula')."""
    for separador in _SEPARADORES:
        if separador in interior:
            clave, _, filtro = interior.partition(separador)
            return clave.strip(), filtro.strip()
    return interior.strip(), ""


def contexto(tipo, datos):
    """Valores ya formateados y listos para sustituir, más las frases derivadas.

    Solo entra lo que el payload trae: una clave ausente tiene que hacer fallar el
    render, no emitir un hueco en un documento legal. Ese mismo criterio vale para una
    frase derivada que no cubra la opción elegida — `tipos.validar` lo impide antes.
    """
    valores = {}
    for campo in tipo["campos"]:
        clave = campo["clave"]
        if clave not in datos:
            continue
        valor = datos[clave]
        formato = FORMATOS.get(campo["tipo"])
        valores[clave] = formato(valor) if formato else "" if valor is None else str(valor)
        # este módulo no sabe nada de género: solo que una opción puede arrastrar frases
        for marcador, por_opcion in campo.get("derivados", {}).items():
            if valor in por_opcion:
                valores[marcador] = por_opcion[valor]
    return valores


def render_markdown(tipo, datos):
    """Renderiza el cuerpo del tipo como Markdown sustituyendo los marcadores."""
    valores = contexto(tipo, datos)

    def sustituir(coincidencia):
        clave, filtro = separar_marcador(coincidencia.group(1))
        if clave not in valores:
            # el equivalente de StrictUndefined: sin esto la clave que falte saldría
            # como un hueco en blanco dentro de un contrato firmado
            raise ValueError(
                f"El cuerpo usa «{{{{{clave}}}}}», que no está en los datos. "
                f"Disponibles: {', '.join(sorted(valores))}"
            )
        if filtro and filtro not in FILTROS:
            raise ValueError(f"El filtro «{filtro}» no existe")
        valor = valores[clave]
        return FILTROS[filtro](valor) if filtro else valor

    return MARCADOR.sub(sustituir, tipo["cuerpo"])


def _configurar_estilos(doc):
    """Calibri 11 con espaciado compacto en el estilo Normal."""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    # la plantilla base de python-docx trae 10 pt de separación y 1,15 de
    # interlineado, demasiado suelto para un documento de cuatro páginas
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.0


def _configurar_pagina(seccion):
    """Carta con márgenes de 1"; deja aire para el encabezado y el pie de 4 líneas."""
    seccion.page_width = Inches(8.5)
    seccion.page_height = Inches(11)
    seccion.left_margin = Inches(1)
    seccion.right_margin = Inches(1)
    seccion.top_margin = Inches(1.3)
    seccion.bottom_margin = Inches(1.1)
    seccion.header_distance = Inches(0.4)
    seccion.footer_distance = Inches(0.35)


def _construir_encabezado(seccion, titulo):
    """Encabezado de todas las páginas: logo a la izquierda y título centrado."""
    parrafo = seccion.header.paragraphs[0]  # el acceso crea la definición del encabezado
    parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ancho = seccion.page_width - seccion.left_margin - seccion.right_margin
    # el tabulador del estilo Header viene fijo en 4680 twips: se recalcula para
    # que siga cayendo en el centro si algún día cambian los márgenes
    parrafo.paragraph_format.tab_stops.add_tab_stop(
        Emu(ancho // 2), WD_TAB_ALIGNMENT.CENTER
    )
    parrafo.add_run().add_picture(str(LOGO), width=Inches(1.2))
    parrafo.add_run("\t")
    parrafo.add_run(titulo).bold = True


def _construir_pie(seccion):
    """Pie de todas las páginas: 4 líneas a 6,5 pt pegadas, la primera en negrita."""
    pie = seccion.footer
    parrafos = [pie.paragraphs[0]] + [pie.add_paragraph() for _ in PIE[1:]]
    for i, (parrafo, texto) in enumerate(zip(parrafos, PIE)):
        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # las 4 líneas van pegadas: el pie hereda de Normal, que sí separa párrafos
        parrafo.paragraph_format.space_after = Pt(0)
        parrafo.paragraph_format.line_spacing = 1.0
        run = parrafo.add_run(texto)
        run.font.size = Pt(6.5)
        run.bold = i == 0


_NEGRITA = re.compile(r"\*\*(.+?)\*\*")
# exige al menos un guion: `|---|---|` es separador de encabezado y se descarta,
# pero `| | |` es una fila real (el espacio en blanco del bloque de firmas)
_SEPARADOR_TABLA = re.compile(r"^\|[\s:|-]*-[\s:|-]*\|$")
_SIN_BORDES = "<!-- tabla-sin-bordes -->"


def _celdas(linea):
    """'| a | b |' -> ['a', 'b']."""
    return [celda.strip() for celda in linea.strip("|").split("|")]


def _bloques(md):
    """Agrupa el Markdown en tuplas (clase, contenido, extra) listas para escribir."""
    bloques = []
    parrafo = []
    filas = []
    sangria = 0
    sin_bordes = False

    def cerrar():
        nonlocal parrafo, filas, sangria, sin_bordes
        if parrafo:
            bloques.append(("parrafo", " ".join(parrafo), sangria))
            parrafo, sangria = [], 0
        if filas:
            bloques.append(("tabla", filas, sin_bordes))
            filas, sin_bordes = [], False

    for cruda in md.splitlines():
        linea = cruda.strip()
        if not linea:
            cerrar()
        elif linea == _SIN_BORDES:
            cerrar()
            sin_bordes = True  # aplica a la siguiente tabla
        elif linea.startswith("|"):
            if parrafo:
                cerrar()  # no reinicia sin_bordes: eso solo pasa al volcar la tabla
            if not _SEPARADOR_TABLA.match(linea):
                filas.append(_celdas(linea))
        elif linea.startswith("- "):
            cerrar()
            bloques.append(("vineta", linea[2:], 0))
        else:
            if filas:
                cerrar()
            if not parrafo:
                sangria = len(cruda) - len(cruda.lstrip())
            parrafo.append(linea)

    cerrar()
    return bloques


def _escribir_runs(parrafo, texto):
    """Vuelca `texto` en `parrafo` partiendo `**negrita**` en runs separados."""
    for i, fragmento in enumerate(_NEGRITA.split(texto)):
        if not fragmento:
            continue
        run = parrafo.add_run(fragmento)
        run.bold = i % 2 == 1  # los índices impares vienen de dentro de los **
    return parrafo


def _escribir_parrafo(doc, texto, estilo=None, alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Agrega un párrafo justificado con la negrita partida en runs."""
    # justificar aquí y no en el estilo Normal: el encabezado, el pie y las celdas
    # heredan de Normal y no deben quedar justificados
    parrafo = doc.add_paragraph(style=estilo)
    parrafo.alignment = alineacion
    return _escribir_runs(parrafo, texto)


def _anchos(columnas, sin_bordes):
    """Reparte las 6,5" útiles: mitades en las firmas, 40/60 en la tabla de datos."""
    if sin_bordes or columnas != 2:
        return [Inches(6.5 / columnas)] * columnas
    return [Inches(2.6), Inches(3.9)]


def _escribir_tabla(doc, filas, sin_bordes):
    """Escribe una tabla; sin estilo hereda Table Normal, que no pinta bordes."""
    columnas = len(filas[0])
    tabla = doc.add_table(rows=len(filas), cols=columnas)
    if not sin_bordes:
        tabla.style = "Table Grid"
    tabla.autofit = False
    anchos = _anchos(columnas, sin_bordes)
    for columna, ancho in zip(tabla.columns, anchos):
        columna.width = ancho
    for fila, textos in zip(tabla.rows, filas):
        for celda, ancho, texto in zip(fila.cells, anchos, textos):
            celda.width = ancho  # w:gridCol y w:tcW se fijan por separado
            _escribir_runs(celda.paragraphs[0], texto)
    return tabla


def markdown_a_docx(md, titulo=TITULO):
    """Convierte el Markdown ya renderizado a un .docx y devuelve los bytes.

    Soporta solo el subconjunto que usa la plantilla: párrafos (líneas seguidas,
    separados por línea en blanco), viñetas `- `, negrita `**...**`, tablas
    `| a | b |` y la directiva `<!-- tabla-sin-bordes -->`. El encabezado y el pie
    no salen del Markdown: se arman aquí, porque el pie lleva un `|` literal que
    el lector de tablas confundiría. Del tipo solo se toma el título: el logo y las
    cuatro líneas del pie son papelería institucional, iguales en cualquier otrosí.
    """
    doc = Document()
    _configurar_estilos(doc)
    seccion = doc.sections[0]
    _configurar_pagina(seccion)
    _construir_encabezado(seccion, titulo)
    _construir_pie(seccion)

    for clase, contenido, extra in _bloques(md):
        if clase == "tabla":
            _escribir_tabla(doc, contenido, extra)
        elif clase == "vineta":
            _escribir_parrafo(doc, contenido, estilo="List Bullet")
        else:
            parrafo = _escribir_parrafo(doc, contenido)
            if extra:  # cada 4 espacios iniciales del bloque valen un cuarto de pulgada
                parrafo.paragraph_format.left_indent = Inches(0.0625 * extra)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _slug(texto, maximo=100):
    """'David Pérez' -> 'david_perez'."""
    sin_tildes = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode()
    return re.sub(r"[^a-z0-9]+", "_", sin_tildes.lower()).strip("_")[:maximo]


def nombre_archivo(tipo, datos):
    """Nombre del .docx: otrosi_teletrabajo_david_perez_20260803.docx.

    Un tipo que no declare campo de nombre o de fecha degrada sin romperse, así que
    el nombre nunca depende de que el tipo tenga campos concretos.
    """
    partes = [tipo.get("prefijo_archivo") or "otrosi"]
    if clave := tipo.get("campo_nombre"):
        partes.append(_slug(str(datos.get(clave) or "")))
    if clave := tipo.get("campo_fecha_archivo"):
        fecha = datos.get(clave)
        if isinstance(fecha, (datetime, date)):
            fecha = f"{fecha:%Y%m%d}"
        partes.append(str(fecha or ""))
    return "_".join(parte for parte in partes if parte) + ".docx"
